"""Genera la guía de estudio en formato Word (.docx) para el primer día en Pillado."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "01-Guia-Estudio-Pillado.docx"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

def H1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def H2(text):
    doc.add_heading(text, level=2)

def H3(text):
    doc.add_heading(text, level=3)

def P(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold

def BULLET(text):
    doc.add_paragraph(text, style="List Bullet")

def NUM(text):
    doc.add_paragraph(text, style="List Number")

def CODE(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)

def TABLE(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
    doc.add_paragraph()

# ============================================================
# PORTADA
# ============================================================
title = doc.add_heading("Guía de estudio — Primer día en Pillado", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Sistema SICOM-ICEO · Flota industrial · Normativa chilena")
r.italic = True
r.font.size = Pt(12)

doc.add_paragraph()
P("Preparado para: Manuel Olivares (admin@pillado.cl)")
P("Alcance: onboarding operacional + técnico en Pillado")
P("Estado datos: sistema en producción con 55 vehículos, 11 contratos reales, abril 2026")

doc.add_page_break()

# ============================================================
# 1. NEGOCIO
# ============================================================
H1("1. El negocio en 30 segundos")

P(
    "Pillado arrienda aproximadamente 55 vehículos industriales (camiones cisterna de "
    "combustible, aljibes de agua, pluma, carrocerías planas, lubrimóviles, camionetas) "
    "a operaciones mineras en Coquimbo y Atacama."
)

P("Clientes reales actualmente cargados en el sistema:", bold=True)
for c in [
    "CMP Romeral", "CM Cenizas Francke", "Boart Longyear Spence",
    "Drilling", "Rentamaq", "Major", "ESMAX", "Orbit", "TPM",
    "San Gerónimo", "ESM",
]:
    BULLET(c)

P("Base operativa central: Taller Pillado — Coquimbo (código FAE-TALLER-CQB).")
P(
    "El ingreso viene del arriendo. Por lo tanto, la métrica que manda es: "
    "cuántos días del mes cada equipo está generando facturación versus detenido "
    "por mantención, falla o falta de demanda.",
    bold=True,
)

# ============================================================
# 2. SICOM-ICEO
# ============================================================
H1("2. SICOM-ICEO: qué es y qué módulos hay")

P(
    "Sistema web (Next.js 15 + Supabase PostgreSQL) que reemplaza la planilla Excel "
    "\"Panel de Control Flota V30\". Tiene 22 páginas. Módulos clave para el día a día:"
)

TABLE(
    ["Módulo", "Para qué sirve"],
    [
        ["Dashboard", "Vista ejecutiva: ICEO, KPIs, alertas, OTs activas"],
        ["Flota", "Estado diario de los 55 vehículos (dimensión técnica + comercial)"],
        ["Reporte Diario", "Snapshot automático por contrato — lo que se envía al cliente"],
        ["Comercial", "Contratos, tarifas, facturación por equipo"],
        ["Órdenes de Trabajo", "Todo trabajo se canaliza con folio OT-YYYYMM-XXXXX"],
        ["Mantenimiento", "Planes PM, calendario preventivo, pautas fabricante"],
        ["Prevención", "SUSPEL/RESPEL, certificaciones, alertas bloqueantes"],
        ["Cumplimiento", "Vencimientos de documentos que bloquean disponibilidad"],
        ["KPI / ICEO", "21 indicadores en 3 áreas + score consolidado 0-100"],
        ["Abastecimiento", "Rutas de despacho de combustible"],
        ["Auditoría", "Log de cada cambio del sistema"],
    ],
)

H2("Concepto clave: estado dual del activo")

P("Cada equipo tiene DOS dimensiones de estado simultáneas:")

P("Estado técnico/operativo (un solo carácter):", bold=True)
TABLE(
    ["Código", "Significado"],
    [
        ["A", "Arrendado"],
        ["D", "Disponible"],
        ["H", "En entrega"],
        ["R", "En recepción"],
        ["M", "Mantención (preventivo)"],
        ["T", "Taller / correctivo"],
        ["F", "Fuera de servicio"],
        ["V", "En venta"],
        ["U", "Uso interno"],
        ["L", "Leasing"],
    ],
)

P("Estado comercial:", bold=True)
for c in [
    "arrendado · disponible · uso_interno · leasing · en_recepcion · en_venta · comprometido",
]:
    BULLET(c)

P(
    "Por qué importa: en el Excel un equipo podía estar marcado \"disponible\" sin "
    "tener realmente la Revisión Técnica al día o sin checklist aprobado. En SICOM un "
    "trigger de base de datos BLOQUEA el paso a disponible si no tiene un checklist "
    "de 55 ítems aprobado y vigente.",
    bold=True,
)

# ============================================================
# 3. OEE
# ============================================================
doc.add_page_break()
H1("3. Cómo calcula el OEE el sistema")

P(
    "OEE = Overall Equipment Effectiveness. El sistema lo calcula por activo y por "
    "flota mediante la función SQL calcular_oee_activo (definida en la migración 25, "
    "archivo 25_flota_oee_normativa.sql línea 354)."
)

H2("3.1 Fórmula maestra")
CODE("OEE = Disponibilidad Mecánica  x  Utilización Operativa  x  Calidad de Servicio")
P(
    "Cada componente se expresa en porcentaje (0 a 100). El OEE final también se "
    "expresa como porcentaje."
)

H2("3.2 Disponibilidad Mecánica")
CODE("Disponibilidad = (Días período  −  Días NO disponibles) / Días período  x  100")
P("Días NO disponibles = días con estado M + T + F.")
P(
    "Nota importante: los días en estado H (en entrega) NO descuentan disponibilidad, "
    "para no castigar logística normal entre arriendos."
)
P("Fuente de datos: tabla estado_diario_flota (una fila por equipo por día).")

H2("3.3 Utilización Operativa")
CODE("Utilización = Horas productivas / Horas disponibles  x  100")
P(
    "Si no hay datos de horas (columnas horas_operativas y horas_disponibles en 0), el "
    "sistema aplica un fallback por días:"
)
CODE("Utilización (fallback) = Días en (A + U + L) / (Días período − Días no disponibles) x 100")
P("Mide si los equipos disponibles realmente se están arrendando.")

H2("3.4 Calidad de Servicio")
CODE("Calidad = (Servicios totales − No conformidades) / Servicios totales  x  100")
P("Servicios totales = días con estado A, U o L (equipo en operación comercial).")
P(
    "No conformidades = filas en la tabla no_conformidades. Categorías consideradas: "
    "entrega fuera de tiempo, entrega incompleta, incumplimiento norma, incidente de "
    "seguridad, contaminación, documentación incompleta, no conformidad ambiental, "
    "repetición de servicio, falla en terreno."
)
P("Si no hubo servicios en el período, la calidad se reporta como 100 por ciento.")

H2("3.5 Clasificación del OEE de flota")
P(
    "La función calcular_oee_flota (agrega todos los activos de un contrato u operación) "
    "entrega esta clasificación final:"
)
TABLE(
    ["OEE promedio flota", "Clasificación"],
    [
        ["≥ 80%", "Clase Mundial"],
        ["64% — 79%", "Bueno"],
        ["50% — 63%", "Aceptable"],
        ["< 50%", "Deficiente"],
    ],
)

H2("3.6 Ejemplo numérico")
P("Equipo en período de 30 días con el siguiente registro:")
for b in [
    "24 días arrendado (estado A)",
    "3 días en mantención preventiva (estado M)",
    "3 días disponible sin arrendar (estado D)",
    "1 no conformidad registrada (entrega fuera de tiempo)",
]:
    BULLET(b)

P("Cálculo:")
CODE(
    "Disponibilidad  = (30 − 3) / 30                         = 90,00 %\n"
    "Utilización     = 24 / (30 − 3)                         = 88,89 %\n"
    "Calidad         = (24 − 1) / 24                         = 95,83 %\n"
    "\n"
    "OEE             = 0,90 x 0,8889 x 0,9583 x 100          = 76,63 %  (Bueno)"
)

# ============================================================
# 4. NORMATIVA
# ============================================================
doc.add_page_break()
H1("4. Normativa chilena aplicable")

P(
    "El sistema genera alertas automáticas bloqueantes a partir de estas normas, "
    "mediante la función SQL fn_ejecutar_verificaciones_normativas:"
)

TABLE(
    ["Norma", "Qué regula", "Alerta que dispara"],
    [
        ["DS 298", "Transporte sustancias peligrosas: antigüedad máx 15 años; SEMEP conductor", "antiguedad_vehiculo, semep_vencido"],
        ["DS 160", "Hermeticidad estanques, inscripción SEC, sensor de fuga", "hermeticidad_vencida, sec_no_vigente, sensor_fuga"],
        ["DS 132", "Seguridad minera: PTS, inducción de faena", "pts_faltante"],
        ["Ley 16.744", "Accidentes del trabajo: reporte dentro de 24 horas", "accidente_no_reportado"],
        ["Ley 21.561", "Jornada conductor: máx 88 hrs espera/mes, conducción continua", "fatiga_conductor"],
        ["Código del Trabajo", "Feriados y descansos", "Validaciones de jornada"],
    ],
)

H2("Certificaciones que caducan y deben renovarse")
for c in [
    "RT — Revisión Técnica",
    "Hermeticidad — tanques de combustible",
    "TC8 / SEC — inscripción SEC",
    "Certificación de gancho — equipos pluma",
    "SEMEP — certificado psicosensométrico del conductor",
    "Inducción de faena minera",
    "Curso sustancias peligrosas",
]:
    BULLET(c)

H2("Regla de oro")
P(
    "Si un equipo o conductor tiene una alerta de severidad critical (prefijo "
    "\"BLOQUEO\" en el título), NO puede operar. El frontend lo marca visualmente y "
    "el trigger de disponibilidad lo impide a nivel de base de datos.",
    bold=True,
)

# ============================================================
# 5. ESTADO INICIAL ALERTAS
# ============================================================
H1("5. Estado inicial de alertas (abril 2026)")

P(
    "Tras cargar los datos reales de Pillado, el sistema arranca con las siguientes "
    "alertas abiertas que deberás revisar en tus primeros días:"
)

TABLE(
    ["Categoría", "Cantidad", "Detalle"],
    [
        ["Certificaciones flota", "10", "RT vencidas o por vencer, hermeticidades, cert. gancho"],
        ["SEMEP conductores", "8", "Los 8 técnicos seed aún sin fecha SEMEP real cargada"],
        ["Disponibilidad", "7", "Equipos sin checklist de verificación vigente"],
        ["Antigüedad flota", "0", "Ningún equipo supera los 15 años por ahora"],
        ["Fatiga conductor", "0", "Sin excesos de horas de espera detectados"],
    ],
)

# ============================================================
# 6. PERFIL
# ============================================================
H1("6. Tu perfil y permisos")

P(
    "Ingresas como admin@pillado.cl con rol administrador. Esto te da acceso total: "
    "ves y editas los 16 módulos del sistema. Los primeros días conviene mantener este "
    "rol para explorar libremente."
)
P(
    "Cuando se creen los 14 usuarios demo (migración 35), podrás loguear como cada rol "
    "(jefe_operaciones, comercial, prevencionista, etc.) para probar cómo ve el sistema "
    "cada perfil sin tocar tu cuenta real."
)

# ============================================================
# 7. GLOSARIO
# ============================================================
H1("7. Glosario mínimo")

TABLE(
    ["Término", "Significado"],
    [
        ["ICEO", "Índice de Cumplimiento y Excelencia Operacional. Score 0-100 que mezcla Área A (Combustibles), Área B (Mantención fijos) y Área C (Mantención móviles)"],
        ["OT", "Orden de Trabajo. Folio único OT-YYYYMM-XXXXX. Todo trabajo se canaliza por acá"],
        ["PM", "Plan de Mantenimiento preventivo. Puede gatillarse por kilometraje o por horómetro"],
        ["SUSPEL", "Sustancias peligrosas"],
        ["RESPEL", "Residuos peligrosos"],
        ["SEMEP", "Certificado psicosensométrico del conductor (DS 298)"],
        ["Faena", "Ubicación física donde opera un contrato (ej. CMP Romeral, CM Francke)"],
        ["Contrato", "Acuerdo comercial con cliente. Cada contrato tiene N faenas y N activos"],
        ["PPU", "Placa Patente Única"],
        ["Sensor ARM / OHP", "Telemetría GPS (requerida por Ley 21.561)"],
        ["RT", "Revisión Técnica vehicular"],
        ["TC8 / SEC", "Superintendencia de Electricidad y Combustibles — inscripción"],
    ],
)

# ============================================================
# 8. CHECKLIST
# ============================================================
doc.add_page_break()
H1("8. Checklist primera semana")

for item in [
    "Navegar los 22 módulos del sistema con datos reales",
    "Entender el flujo de una OT extremo a extremo (crear, ejecutar, cerrar)",
    "Leer la matriz de permisos por rol (tabla _roles_matriz_permisos)",
    "Visitar el Taller Pillado Coquimbo y conocer a los 8 técnicos del seed",
    "Revisar las 25 alertas abiertas y categorizarlas por urgencia",
    "Crear los 14 usuarios demo en Supabase y ejecutar migración 35",
    "Leer documentos docs/ARQUITECTURA-OBJETIVO.md y docs/MANUAL-KPI-COMPLETO.md",
    "Proponer 2 mejoras al gerente general basándose en lo observado",
]:
    BULLET(item)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Documento de estudio — uso interno")
r.italic = True
r.font.size = Pt(9)

doc.save(OUT)
print(f"OK: {OUT} generado")
