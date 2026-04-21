"""Genera el Word de estudio con el resumen operacional detallado de Pillado."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "05-Resumen-Operacion-Pillado.docx"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def H1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def H2(text):
    doc.add_heading(text, level=2)


def H3(text):
    doc.add_heading(text, level=3)


def P(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def BULLET(text):
    doc.add_paragraph(text, style="List Bullet")


def NUM(text):
    doc.add_paragraph(text, style="List Number")


def CODE(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)


def TABLE(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
    doc.add_paragraph()


def CALLOUT(text):
    p = doc.add_paragraph()
    run = p.add_run("⚠  " + text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xB0, 0x3A, 0x2E)


# Helper para tablas de equipos por contrato
EQUIPO_HEADERS = ["Patente", "Código interno", "Marca y modelo", "Año", "Antig.", "Capacidad / tipo"]

# ============================================================
# PORTADA
# ============================================================
title = doc.add_heading("Operación Pillado — Estado actual cargado en SICOM", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Contratos · Flota detallada · Operadores · Certificaciones — Abril 2026")
r.italic = True
r.font.size = Pt(12)

doc.add_paragraph()
P("Preparado para: Manuel Olivares (admin@pillado.cl)")
P("Fuente: seeds reales de las migraciones 26, 28, 29 aplicadas en Supabase")
P("Alcance: patente por patente, por contrato, por faena. Todo lo que está cargado hoy.")

doc.add_paragraph()
P("Cómo leer este documento", bold=True)
P(
    "Cada contrato trae la tabla completa de sus equipos con patente, marca+modelo, "
    "año y capacidad. Al final hay una sección con los 21 equipos de flota interna "
    "(sin contrato), los operadores y las acciones críticas para la primera semana."
)

doc.add_page_break()

# ============================================================
# 1. RESUMEN EJECUTIVO
# ============================================================
H1("1. Resumen ejecutivo")

TABLE(
    ["Métrica", "Valor"],
    [
        ["Total vehículos activos en la flota", "55"],
        ["Contratos activos", "11"],
        ["Equipos bajo contrato", "34"],
        ["Equipos flota interna (sin contrato / taller)", "21"],
        ["Contrato más grande", "Boart Longyear (11 equipos)"],
        ["Equipos sobre 15 años (bloqueo DS 298)", "1 (DCHD-83, 2011)"],
        ["Edad promedio de la flota", "≈ 6-7 años"],
        ["Operadores/técnicos cargados", "8"],
        ["Faenas principales", "Calama, Coquimbo, Taltal, Copiapó, El Salvador, Los Bronces"],
    ],
)

# ============================================================
# 2. CONTRATOS Y EQUIPOS POR FAENA
# ============================================================
H1("2. Contratos con detalle de equipos por faena")

# --- BOART ---
H2("2.1 Boart Longyear — CTR-BOART-2024 (11 equipos)")
P("Faenas: Spence y División Ministro Hales, Calama. Vigencia 2024-06-01 a 2027-05-31.")
P("Es el contrato más grande de Pillado (20 % de la flota). Mezcla cisternas, plumas, polibrazo y carrocerías.", bold=True)
TABLE(
    EQUIPO_HEADERS,
    [
        ["TGGF-56", "AI-22-01", "Volvo FMX 420",  "2024", "2", "Cisterna 22 kL"],
        ["TGGF-57", "AI-22-02", "Volvo FMX 420",  "2024", "2", "Cisterna 22 kL"],
        ["TGGF-58", "AI-22-03", "Volvo FMX 420",  "2024", "2", "Cisterna 22 kL"],
        ["TGGF-59", "CH-20-01", "Volvo FMX 420",  "2024", "2", "Polibrazo 20 t"],
        ["TRDP-96", "CP-06-03", "Volvo FMX 420",  "2024", "2", "Pluma 10 t"],
        ["TRST-57", "CC-20-03", "Scania P450B",   "2025", "1", "Cisterna 20 kL"],
        ["TRST-58", "AI-20-06", "Scania P450B",   "2025", "1", "Cisterna 20 kL"],
        ["TRSS-13", "CS-06-03", "Scania P450B",   "2025", "1", "Carrocería plana 14.5 t"],
        ["TRSS-14", "CP-06-04", "Scania P450B",   "2025", "1", "Pluma 10 t"],
        ["TRSS-15", "CS-06-04", "Scania P450B",   "2025", "1", "Carrocería plana 14.5 t"],
        ["TTPC-47", "CS-06-05", "Scania P450B",   "2025", "1", "Carrocería plana 14.5 t"],
    ],
)

# --- CMP ---
H2("2.2 CMP Romeral — CTR-CMP-2025 (4 equipos)")
P("Faena: Romeral (Coquimbo). Vigencia 2025-01-01 a 2027-12-31. Uso interno combustibles.")
TABLE(
    EQUIPO_HEADERS,
    [
        ["DJKL-18", "CC-15-05", "Mercedes-Benz Actros 3341", "2012", "14", "Cisterna comb. 15 kL"],
        ["FSLZ-67", "CC-15-06", "Mercedes-Benz Actros 3341", "2013", "13", "Cisterna comb. 15 kL"],
        ["LCSX-78", "CC-15-17", "Mack GU 813 aut.",          "2019", "7",  "Cisterna comb. 15 kL"],
        ["RZPC-83", "CA-28-02", "Maxus T60 4x4 DX",          "2022", "4",  "Camioneta soporte"],
    ],
)
CALLOUT("LCSX-78 tiene hermeticidad vencida — no puede cargar combustible hasta renovarla.")

# --- CENIZAS ---
H2("2.3 CM Cenizas Francke — CTR-CENIZAS-2025 (3 equipos)")
P("Faena: Taltal (Antofagasta). Vigencia 2025-01-01 a 2027-12-31. Administración combustibles.")
TABLE(
    EQUIPO_HEADERS,
    [
        ["HHWB-42", "CC-20-01", "Mack GU 813 aut.",           "2015", "11", "Cisterna comb. 20 kL"],
        ["HHWB-44", "CC-20-02", "Mack GU 813 aut.",           "2015", "11", "Cisterna comb. 20 kL"],
        ["LLBP-96", "CA-28-01", "Toyota Hilux 2.8 Autom",     "2019", "7",  "Camioneta soporte"],
    ],
)
CALLOUT("HHWB-42 y HHWB-44 tienen hermeticidad vencida. Equipos de Pillado más antiguos en Taltal.")

# --- DRILLING ---
H2("2.4 Drilling Service and Solution — CTR-DRILLING-2025 (4 equipos)")
P("Faenas: Sobek y Cuprita (Copiapó), Huachalalume (La Serena). Vigencia 2025-01-01 a 2026-12-31.")
TABLE(
    EQUIPO_HEADERS,
    [
        ["JTYK-88", "AI-20-02", "Mercedes-Benz Actros 3336 K",  "2017", "9", "Cisterna riego 20 kL"],
        ["JGBY-10", "CC-15-13", "Mercedes-Benz Axor 2633/45",   "2017", "9", "Cisterna comb. 15 kL"],
        ["KCBY-30", "AI-25-02", "Mercedes-Benz Actros 3336 K",  "2018", "8", "Cisterna riego 25 kL"],
        ["KVWW-69", "AI-25-05", "Mercedes-Benz Actros 3336 K",  "2019", "7", "Cisterna riego 25 kL"],
    ],
)

# --- RENTAMAQ ---
H2("2.5 Rentamaq → Teck Andacollo — CTR-RENTAMAQ-2025 (2 equipos)")
P("Faena: Teck Andacollo. Vigencia 2025-01-01 a 2026-06-30 (renegociar primer semestre).")
TABLE(
    EQUIPO_HEADERS,
    [
        ["SVCZ-38", "AI-20-04", "Volvo VM 350", "2023", "3", "Cisterna riego 20 kL"],
        ["SVBJ-55", "AI-20-05", "Volvo VM 350", "2023", "3", "Cisterna riego 20 kL"],
    ],
)

# --- ORBIT ---
H2("2.6 Orbit Garant — CTR-ORBIT-2025 (2 equipos)")
P("Faenas: Los Bronces (Santiago) y El Abra (Calama). Vigencia 2025-01-01 a 2026-12-31.")
TABLE(
    EQUIPO_HEADERS,
    [
        ["SVBJ-57", "CC-15-15", "Volvo VM 350", "2023", "3", "Cisterna comb. 15 kL"],
        ["TCJV-15", "CC-15-16", "Renault C440", "2024", "2", "Cisterna comb. 15 kL"],
    ],
)
CALLOUT("TCJV-15 tiene hermeticidad vencida. SVBJ-57 tiene RT por vencer en 32 días.")

# --- ESM ---
H2("2.7 ESM Calama — CTR-ESM-2025 (2 equipos)")
P("Faena: Calama. Vigencia 2025-01-01 a 2026-12-31. Servicios de vehículos.")
TABLE(
    EQUIPO_HEADERS,
    [
        ["SPRY-26", "CA-24-02", "Toyota Hilux 4x4 2.4 MT DX", "2023", "3", "Camioneta"],
        ["SPRY-28", "CA-24-03", "Toyota Hilux 4x4 2.4 MT DX", "2023", "3", "Camioneta"],
    ],
)

# --- ESMAX ---
H2("2.8 Esmax El Salvador — CTR-ESMAX-2025 (1 equipo)")
P("Faena: El Salvador (CODELCO). Vigencia 2025-01-01 a 2026-06-30.")
TABLE(
    EQUIPO_HEADERS,
    [
        ["SVBJ-56", "CC-15-14", "Volvo VM 350", "2023", "3", "Cisterna comb. 15 kL"],
    ],
)

# --- SAN GERONIMO ---
H2("2.9 San Gerónimo — CTR-SANGERONIMO-2025 (2 equipos)")
P("Faena: Lambert. Vigencia 2025-01-01 a 2026-12-31. Aljibes de combustible.")
TABLE(
    EQUIPO_HEADERS,
    [
        ["KVWD-27", "CC-05-10", "Mercedes-Benz Accelo 1016/44",  "2018", "8",  "Cisterna comb. 5 kL"],
        ["FJTJ-60", "CC-44-03", "Mercedes-Benz Atego 1624A 4x4", "2013", "13", "Cisterna comb. 5 kL 4x4"],
    ],
)
CALLOUT("FJTJ-60 tiene hermeticidad vencida y es el segundo equipo más antiguo del contrato (13 años).")

# --- MAJOR ---
H2("2.10 Major Drilling Yastai — CTR-MAJOR-2025 (1 equipo)")
P("Faena: Tierra Amarilla. Vigencia 2025-06-01 a 2026-05-31.")
TABLE(
    EQUIPO_HEADERS,
    [
        ["RSCY-85", "CC-05-11", "Mercedes-Benz Accelo 1016/44", "2022", "4", "Cisterna comb. 5 kL"],
    ],
)

# --- TPM ---
H2("2.11 TPM Minería Caserones — CTR-TPM-2025 (1 equipo)")
P("Faena: Caserones (Copiapó). Vigencia 2025-01-01 a 2026-12-31. Camioneta lubricadora.")
TABLE(
    EQUIPO_HEADERS,
    [
        ["SBPG-12", "CA-24-01", "Toyota Hilux 4x4 2.4 MT DX", "2022", "4", "Camioneta lubricadora"],
    ],
)

# ============================================================
# 3. FLOTA INTERNA SIN CONTRATO
# ============================================================
doc.add_page_break()
H1("3. Flota interna sin contrato asignado (21 equipos)")

P(
    "Son los equipos disponibles en el taller Coquimbo / Calama para uso interno, "
    "flota de gerencia, apoyo a operaciones o para asignar a nuevos contratos. "
    "Representan un 38 % de la flota total — potencial comercial."
)

H2("3.1 Cisternas y equipos pesados disponibles")
TABLE(
    EQUIPO_HEADERS,
    [
        ["TRDP-97", "AI-22-04", "Volvo FMX 420",              "2024", "2",  "Cisterna 22 kL"],
        ["GCHT-12", "AI-25-01", "Mack GU813E mec.",            "2014", "12", "Cisterna 25 kL"],
        ["KCBY-31", "AI-25-03", "Mercedes-Benz Actros 3336 K", "2018", "8",  "Cisterna 25 kL"],
        ["KVWW-68", "AI-25-04", "Mercedes-Benz Actros 3336 K", "2019", "7",  "Cisterna 25 kL"],
        ["GGHB-32", "AI-20-03", "Mack GU813E Allison",         "2014", "12", "Cisterna 20 kL"],
        ["LKPY-18", "AI-20-07", "Mack GU 813 aut.",            "2019", "7",  "Cisterna 20 kL"],
        ["HKSR-81", "CC-15-09", "Mercedes-Benz Axor 2633",     "2016", "10", "Cisterna comb. 15 kL"],
        ["DCHD-83", "CC-05-04", "Mitsubishi Canter 7.5",       "2011", "15", "Cisterna comb. 5 kL"],
        ["FJTJ-61", "CC-44-04", "Mercedes-Benz Atego 1624A 4x4","2013", "13", "Chasis cabinado"],
        ["RSCY-86", "CS-06-02", "Mercedes-Benz Accelo 1016/44","2022", "4",  "Carrocería 6 t"],
        ["TGGF-60", "CP-06-02", "Volvo FMX 540",               "2024", "2",  "Pluma 10 t"],
    ],
)
CALLOUT(
    "DCHD-83 (Mitsubishi Canter 2011) está en el límite del DS 298 (15 años). Decidir "
    "entre venta, uso interno no-combustible o baja definitiva antes de fin de año."
)

H2("3.2 Camionetas y equipos livianos")
TABLE(
    EQUIPO_HEADERS,
    [
        ["VRST-19", "CA-20-01", "Maxus T60 4x4 DX Plus 6 MT",       "2025", "1", "Camioneta"],
        ["TSTB-48", "CA-12-01", "Chevrolet Montana 1.2 MT",         "2025", "1", "Camioneta"],
        ["TCRB-71", "OC-GE-02", "RAM 1500 LIMITED 5.7L",            "2024", "2", "Camioneta Gerencia"],
        ["KVDK-20", "CA-23-04", "Nissan NP300 Doble Cabina",        "2019", "7", "Camioneta"],
        ["KVDK-21", "CA-23-05", "Nissan NP300 Doble Cabina",        "2019", "7", "Camioneta"],
        ["JDKH-31", "CA-23-03", "Nissan NP300 Doble Cabina",        "2017", "9", "Camioneta"],
        ["SLRK-82", "VC-FC-02", "Citroën Berlingo K9 1.6 Diesel",   "2023", "3", "Furgón TM-11"],
        ["SPRY-29", "VC-FC-03", "Citroën Berlingo K9 1.6 Diesel",   "2023", "3", "Furgón TM-12"],
    ],
)

H2("3.3 Equipos menores (grúas horquilla del taller)")
TABLE(
    EQUIPO_HEADERS,
    [
        ["GCSY-66",  "GH-05-01", "Toyota 02-7FDA50", "2014", "12", "Grúa horquilla 7.3 t"],
        ["GDP 30TK", "GH-03-01", "Yale GDP 30TK",    "2014", "12", "Grúa horquilla 3 t"],
    ],
)

# ============================================================
# 4. OPERADORES Y TÉCNICOS
# ============================================================
doc.add_page_break()
H1("4. Operadores y técnicos cargados")

H2("4.1 Los 8 técnicos/conductores del taller")
TABLE(
    ["#", "Nombre", "RUT (seed)", "Licencia", "Cargo", "Faena", "SEMEP"],
    [
        ["1", "Felipe López",    "11111111-1", "A2", "Téc. Mecánico Senior", "FAE-TALLER-CQB", "Activo"],
        ["2", "Juan Valenzuela", "22222222-2", "A2", "Técnico Mecánico",      "FAE-TALLER-CQB", "Activo"],
        ["3", "Yohan Rondón",    "33333333-3", "A2", "Técnico Mecánico",      "FAE-TALLER-CQB", "Activo"],
        ["4", "Pereira",         "44444444-4", "A2", "Técnico Mecánico",      "(sin asignar)",  "Activo"],
        ["5", "Luis Hernández",  "55555555-5", "A2", "Técnico Mecánico",      "(sin asignar)",  "Activo"],
        ["6", "Rodrigo Cortés",  "66666666-6", "A2", "Técnico Mecánico",      "(sin asignar)",  "Activo"],
        ["7", "Jesús Varela",    "77777777-7", "B",  "Técnico Mecánico",      "(sin asignar)",  "Activo"],
        ["8", "Nibaldo",         "88888888-8", "B",  "Técnico Mecánico",      "(sin asignar)",  "Activo"],
    ],
)

CALLOUT(
    "Todos los SEMEP están marcados activos pero SIN fecha real cargada. Por eso el "
    "sistema dispara 8 alertas automáticas de tipo 'semep_vencido'. Tarea de día uno: "
    "pedir los certificados reales y hacer UPDATE conductores SET semep_vencimiento = ...;"
)

H2("4.2 Supervisores cargados en usuarios_perfil")
TABLE(
    ["Cargo", "Faena", "Notas"],
    [
        ["Supervisor de Operaciones Coquimbo", "FAE-TALLER-CQB", "Responsable de la zona Coquimbo/Romeral/Andacollo"],
        ["Supervisor de Operaciones Calama",   "FAE-TALLER-CAL", "Responsable de la zona Calama/Spence/Ministro Hales"],
        ["Administración de contrato",         "(centralizado)",  "Sin faena fija — ve todos los contratos"],
    ],
)

H2("4.3 Carga de trabajo actual — estado al momento del seed")

P(
    "Los seeds 26, 28 y 29 NO crean órdenes de trabajo abiertas. Tampoco asignan un "
    "conductor titular a cada activo (la tabla activos no tiene columna conductor_id). "
    "Por lo tanto la carga de trabajo de cada técnico es 0 en el snapshot inicial. "
    "La carga real nacerá de:"
)
for i in [
    "OTs que tú crees al marcar equipos M o T en /dashboard/flota.",
    "Planes preventivos que venzan por kilometraje u horas (fn_verificar_planes_preventivos).",
    "Certificaciones vencidas que requieren OT de mantención (5 hermeticidades, 1 RT, 3 ganchos urgentes).",
    "Asignación manual de operadores titulares por equipo cuando decidas la política.",
]:
    BULLET(i)

P("Propuesta operacional para cubrir la brecha de asignación:", bold=True)
for s in [
    "Asignar Felipe López como líder técnico del taller Coquimbo (8 años de experiencia implícita en el cargo Senior).",
    "Juan Valenzuela y Yohan Rondón como ejecutores día a día de OTs preventivas y correctivas en Coquimbo.",
    "Rodrigo Cortés o Pereira a cargo de salidas a terreno en contratos remotos (Taltal, Andacollo).",
    "Jesús Varela y Nibaldo (licencia B) para mantención en taller sin salida a ruta.",
]:
    BULLET(s)
P(
    "Esta asignación debe conversarse con el gerente general y el supervisor de "
    "operaciones antes de dejarla fija en el sistema."
)

# ============================================================
# 5. CERTIFICACIONES CRÍTICAS
# ============================================================
H1("5. Certificaciones críticas al inicio")

H2("5.1 Vencidas — acción inmediata")
TABLE(
    ["Patente", "Contrato", "Tipo certificación", "Acción"],
    [
        ["TRSS-13", "Boart Longyear",    "Revisión Técnica vencida 2026-04-08", "Agendar RT urgente"],
        ["TCJV-15", "Orbit Garant",      "Hermeticidad vencida",                 "Prueba hidrostática"],
        ["LCSX-78", "CMP Romeral",       "Hermeticidad vencida",                 "Prueba hidrostática"],
        ["HHWB-42", "CM Cenizas Francke","Hermeticidad vencida",                 "Prueba hidrostática"],
        ["HHWB-44", "CM Cenizas Francke","Hermeticidad vencida",                 "Prueba hidrostática"],
        ["FJTJ-60", "San Gerónimo",      "Hermeticidad vencida",                 "Prueba hidrostática"],
        ["TGGF-60", "Flota interna",     "Cert. gancho vencido",                 "Renovar organismo certif."],
        ["TRDP-96", "Boart Longyear",    "Cert. gancho vencido",                 "Renovar organismo certif."],
        ["TRSS-16", "Boart Longyear",    "Cert. gancho vencido",                 "Renovar organismo certif."],
    ],
)

H2("5.2 Por vencer en menos de 45 días")
TABLE(
    ["Patente", "Contrato", "Tipo", "Días restantes"],
    [
        ["SVCZ-38", "Rentamaq/Teck",   "RT", "27"],
        ["SVBJ-57", "Orbit Garant",    "RT", "32"],
        ["JDKH-31", "Flota interna",   "RT", "20"],
        ["KVDK-21", "Flota interna",   "RT", "20"],
        ["TCRB-71", "Flota interna",   "RT", "20"],
    ],
)

CALLOUT(
    "5 de los 16 aljibes de combustible tienen hermeticidad vencida. Hasta renovarlas "
    "esos equipos NO pueden transportar combustible legalmente, aunque la UI los "
    "muestre como 'arrendados'. Es el tema más urgente de Pillado hoy."
)

# ============================================================
# 6. MÉTRICAS CLAVE POR CONTRATO Y SQL ÚTIL
# ============================================================
doc.add_page_break()
H1("6. Métricas por contrato y SQL de consulta")

P(
    "Hoy no hay snapshot persistente segmentado por contrato. Estas consultas te "
    "responden las preguntas más frecuentes que te hará el gerente:"
)

H2("6.1 Disponibilidad del día por contrato")
CODE(
    "SELECT c.codigo, c.nombre,\n"
    "       COUNT(*) FILTER (WHERE e.estado_codigo = 'A')        AS arrendados,\n"
    "       COUNT(*) FILTER (WHERE e.estado_codigo = 'D')        AS disponibles,\n"
    "       COUNT(*) FILTER (WHERE e.estado_codigo IN ('M','T')) AS en_taller,\n"
    "       COUNT(*) FILTER (WHERE e.estado_codigo = 'F')        AS fuera_servicio\n"
    "FROM contratos c\n"
    "JOIN activos a             ON a.contrato_id = c.id\n"
    "JOIN estado_diario_flota e ON e.activo_id   = a.id\n"
    "                           AND e.fecha      = CURRENT_DATE\n"
    "GROUP BY c.codigo, c.nombre\n"
    "ORDER BY COUNT(*) DESC;"
)

H2("6.2 OEE mensual por operación")
CODE(
    "SELECT * FROM calcular_oee_flota(\n"
    "    NULL,\n"
    "    date_trunc('month', CURRENT_DATE)::DATE,\n"
    "    CURRENT_DATE,\n"
    "    'Calama'   -- o 'Coquimbo'\n"
    ");"
)

H2("6.3 Todas las certificaciones vencidas o por vencer 45 días")
CODE(
    "SELECT a.patente, c.codigo AS contrato, ce.tipo, ce.fecha_vencimiento,\n"
    "       (ce.fecha_vencimiento - CURRENT_DATE) AS dias_restantes\n"
    "FROM certificaciones ce\n"
    "JOIN activos a    ON a.id = ce.activo_id\n"
    "LEFT JOIN contratos c ON c.id = a.contrato_id\n"
    "WHERE ce.bloqueante = true\n"
    "  AND ce.fecha_vencimiento <= CURRENT_DATE + INTERVAL '45 days'\n"
    "ORDER BY ce.fecha_vencimiento;"
)

H2("6.4 Equipos por contrato y faena (inventario completo)")
CODE(
    "SELECT c.codigo, a.patente, a.codigo AS cod_interno,\n"
    "       m.descripcion AS modelo, a.anio_fabricacion AS anio,\n"
    "       a.ubicacion_actual AS faena\n"
    "FROM activos a\n"
    "LEFT JOIN contratos c ON c.id = a.contrato_id\n"
    "LEFT JOIN modelos   m ON m.id = a.modelo_id\n"
    "WHERE a.estado != 'dado_baja'\n"
    "ORDER BY c.codigo NULLS LAST, a.patente;"
)

H2("6.5 SEMEP de conductores")
CODE(
    "SELECT nombre_completo, rut, semep_vigente, semep_vencimiento, semep_tipo\n"
    "FROM conductores\n"
    "WHERE activo = true\n"
    "ORDER BY semep_vencimiento NULLS FIRST;"
)

# ============================================================
# 7. PRIORIZACIÓN PRIMERA SEMANA
# ============================================================
H1("7. Plan sugerido — primera semana")

NUM("Visitar el taller Coquimbo, reunirse con los 8 técnicos y confirmar cargos y faenas.")
NUM("Pedir los 8 SEMEP reales y hacer UPDATE en tabla conductores (apaga 8 alertas).")
NUM("Agendar prueba de hermeticidad de TCJV-15, LCSX-78, HHWB-42, HHWB-44, FJTJ-60.")
NUM("Agendar RT de TRSS-13 (vencida) y de los 5 equipos con RT a <45 días.")
NUM("Decidir destino de DCHD-83 (Mitsubishi Canter 2011, 15 años — límite DS 298).")
NUM("Renovar certificados de gancho de TGGF-60, TRDP-96, TRSS-16 (los 3 plumas bloqueados).")
NUM("Reunión con Boart Longyear — contrato más grande (11 equipos) y en Calama. Validar estado de los 11 equipos en terreno.")
NUM("Definir política de asignación operador→equipo y cargar en sistema.")
NUM("Presentar al gerente general el mapa completo de alertas y priorización.")

# ============================================================
# 8. GLOSARIO DE CLIENTES
# ============================================================
H1("8. Glosario de clientes y contratos")

TABLE(
    ["Cliente", "Quiénes son", "Peso para Pillado"],
    [
        ["Boart Longyear",   "Multinacional de servicios de perforación",       "Contrato más grande (11 equipos, 20 % flota)"],
        ["CMP",              "Compañía Minera del Pacífico — hierro (Romeral)", "Contrato histórico Coquimbo"],
        ["CM Cenizas",       "Minera polimetálica en Taltal",                   "Operación remota y calurosa"],
        ["Drilling S&S",     "Servicios de perforación menor",                  "Mezcla aljibes riego + combustible"],
        ["Rentamaq/Teck",    "Arriendo de maquinaria → Teck Andacollo",          "Teck es de los clientes mineros más exigentes"],
        ["Orbit Garant",     "Perforación exploración",                         "Dos aljibes en Santiago + El Abra"],
        ["ESM",              "Empresa servicios minería — Calama",              "Camionetas soporte"],
        ["Esmax",            "Distribuidora de combustibles",                    "Aljibe combustible El Salvador (CODELCO)"],
        ["San Gerónimo",     "Minera polimetálica",                              "Aljibes combustible Lambert"],
        ["Major Drilling",   "Perforación global",                                "Aljibe 5 kL Tierra Amarilla"],
        ["TPM",              "Contratista Caserones (Lumina Copper)",           "Camioneta lubricadora Caserones"],
    ],
)

# ============================================================
# CIERRE
# ============================================================
doc.add_paragraph()
H1("Notas de cierre")

P(
    "Esta foto es del seed aplicado (abril 2026). Los datos cambian en cuanto el "
    "sistema opera: el modal de /dashboard/flota mueve el estado diario, el cron de "
    "06:30 regenera el snapshot, fn_ejecutar_verificaciones_normativas() redescubre "
    "alertas. Usa este documento como baseline, no como verdad permanente."
)
P(
    "En 30 días, el reporte diario con tendencia (migración 36) ya te habrá mostrado "
    "patrones reales: qué contratos bajan disponibilidad, qué equipos entran seguido "
    "a taller, qué técnico cierra más OTs. Ese será el siguiente nivel de análisis."
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Documento de estudio — uso interno Pillado")
r.italic = True
r.font.size = Pt(9)

doc.save(OUT)
print(f"OK: {OUT} generado")
