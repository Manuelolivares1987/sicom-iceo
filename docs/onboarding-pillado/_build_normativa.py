"""Genera la guía de estudio de NORMATIVA chilena aplicable a Pillado (.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "03-Guia-Normativa-Pillado.docx"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def H1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def H2(text):
    doc.add_heading(text, level=2)


def H3(text):
    doc.add_heading(text, level=3)


def P(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def BULLET(text):
    doc.add_paragraph(text, style="List Bullet")


def NUM(text):
    doc.add_paragraph(text, style="List Number")


def CODE(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)


def TABLE(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
    doc.add_paragraph()


def CALLOUT(text):
    p = doc.add_paragraph()
    run = p.add_run("⚠  " + text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xB0, 0x3A, 0x2E)


# ============================================================
# PORTADA
# ============================================================
title = doc.add_heading("Guía de normativa chilena aplicable a Pillado", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Transporte de cargas peligrosas · Seguridad laboral · Medio ambiente")
r.italic = True
r.font.size = Pt(12)

doc.add_paragraph()
P("Preparado para: Manuel Olivares (admin@pillado.cl)")
P("Alcance: material de estudio operacional para la flota de 55 vehículos industriales")
P("Operaciones: Coquimbo y Calama/Atacama — minería")

doc.add_paragraph()
P("Cómo leer esta guía", bold=True)
P(
    "Las normas están ordenadas por utilidad práctica, NO por número de decreto. "
    "Primero lo que te pide el cliente minero cada mañana (vehículo + conductor), "
    "después lo laboral/prevención de riesgos, y al final lo ambiental (sustancias "
    "y residuos peligrosos)."
)
P(
    "Cada norma tiene: qué regula, obligaciones concretas, documentos que debes "
    "pedir o renovar, y cómo se refleja en el sistema SICOM-ICEO."
)

doc.add_page_break()

# ============================================================
# BLOQUE A — VEHÍCULO Y CONDUCTOR
# ============================================================
H1("Bloque A — Vehículo y conductor (transporte)")

# --------- DS 298 ---------
H2("A.1  DS 298/1994 MTT — Transporte de cargas peligrosas")

P("Qué regula", bold=True)
P(
    "Condiciones técnicas que deben cumplir los camiones (y sus conductores) que "
    "transportan sustancias peligrosas por caminos públicos. Aplica a cisternas de "
    "combustible, aljibes con productos químicos, transporte de ácidos, oxidantes, "
    "inflamables y corrosivos. Es la norma más fiscalizada en terreno por "
    "Carabineros (sección SIAT) y el MTT."
)

P("Obligaciones para Pillado", bold=True)

P("Sobre el vehículo:", bold=True)
BULLET("Antigüedad máxima de 15 años para transportar cargas peligrosas.")
BULLET("Rotulado reglamentario: rombos de peligro ONU visibles en ambos costados y parte trasera, más número ONU de 4 dígitos bajo el rombo.")
BULLET("Equipamiento de emergencia: 2 extintores mínimo (uno en cabina, otro accesible desde afuera), triángulos reflectantes, calzos, botiquín.")
BULLET("Sistema de corte de corriente accesible y señalizado.")
BULLET("Conexión a tierra (pértiga) obligatoria para cisternas de combustible durante carga y descarga.")

P("Sobre el conductor:", bold=True)
BULLET("SEMEP — examen psicosensométrico obligatorio en institución autorizada.")
BULLET("Renovación: anual para conductores de vehículos pesados y cargas peligrosas; cada 4 años para licencias livianas.")
BULLET("Licencia profesional A2, A3, A4 o A5 según tipo de vehículo.")
BULLET("Curso aprobado de transporte de sustancias peligrosas (vigente).")
BULLET("Hoja de ruta con datos del producto, cantidad, destino, teléfono de emergencia.")

P("Documentos que debes pedir y mantener vigentes", bold=True)
for d in [
    "Certificado SEMEP por cada conductor (los 8 técnicos del seed aún sin fecha real → 8 alertas abiertas)",
    "Copia de licencia profesional de cada conductor",
    "Certificado curso transporte sustancias peligrosas",
    "Fotos del rotulado actual de cada camión cisterna",
    "Registro de entrega de EPP al conductor",
]:
    BULLET(d)

P("Cómo lo controla el sistema", bold=True)
CODE("Tabla conductores  →  columnas semep_vigente, semep_vencimiento, semep_tipo")
CODE("Función fn_verificar_semep_conductores()  →  alerta 'semep_vencido'")
CODE("Función fn_verificar_antiguedad_flota()   →  alerta 'antiguedad_vehiculo'")

CALLOUT(
    "Un conductor con SEMEP vencido NO puede conducir cargas peligrosas. "
    "En Pillado esto se traduce en que no puede operar ninguna cisterna de combustible."
)

# --------- DS 160 ---------
H2("A.2  DS 160/2008 Economía — Combustibles líquidos")

P("Qué regula", bold=True)
P(
    "Seguridad en el transporte y almacenamiento de combustibles líquidos (diesel, "
    "bencina, kerosene, petróleo). Es la base que exige la SEC (Superintendencia de "
    "Electricidad y Combustibles) para autorizar la operación de cisternas."
)

P("Obligaciones sobre la cisterna", bold=True)
BULLET("Inscripción en la SEC de cada camión cisterna (sin esto, NO puede circular con combustible).")
BULLET("Certificado TC8 — aptitud técnica emitido por organismo de certificación autorizado (CESMEC, DICTUC, DIMAQ). Renovación anual.")
BULLET("Prueba de hermeticidad anual del estanque (hidrostática + neumática).")
BULLET("Válvulas de seguridad: válvula de fondo + válvula de emergencia activable desde cabina.")
BULLET("Sensor de fuga y sistema de detección de vapores en bombas.")
BULLET("Manhole (bocas de carga) con cierre hermético y sello funcional.")
BULLET("Compartimentos separados si transporta más de un producto.")

P("Obligaciones sobre el operador", bold=True)
BULLET("Curso de operación segura de cisternas aprobado.")
BULLET("EPP específico: traje antiestático, calzado conductivo, guantes dieléctricos en instalaciones eléctricas.")
BULLET("Procedimiento documentado de carga y descarga con puesta a tierra.")

P("Documentos que debes pedir", bold=True)
for d in [
    "TC8 vigente de cada camión cisterna",
    "Certificado de hermeticidad vigente",
    "Inscripción SEC activa",
    "Permisos de transporte de combustibles por región",
    "Registros de las últimas 3 pruebas de hermeticidad",
]:
    BULLET(d)

P("Cómo lo controla el sistema", bold=True)
CODE("Tabla certificaciones  →  tipo = 'tc8_sec' | 'hermeticidad' | 'inscripcion_sec'")
CODE("Función fn_verificar_certificaciones_flota()")
P("Alertas: tc8_por_vencer, hermeticidad_vencida, sec_no_vigente.")

# --------- LEY 21.561 ---------
H2("A.3  Ley 21.561 — Jornada y descansos (\"Ley de las 40 horas\")")

P("Qué regula", bold=True)
P(
    "Jornada máxima laboral, tiempos de descanso obligatorios y fatiga del "
    "conductor. Es la norma que reemplazó y endureció la antigua regulación sobre "
    "jornada en transporte. Está en plena implementación por etapas hasta 2028."
)

P("Obligaciones clave para conductores de Pillado", bold=True)
BULLET("Máximo 5 horas de conducción continua sin descanso.")
BULLET("Descanso mínimo entre bloques de conducción: 30 minutos cada 5 horas continuas.")
BULLET("Máximo 88 horas de tiempo de espera al mes (tiempo en que el conductor está disponible en faena pero no conduciendo).")
BULLET("Descanso diario obligatorio: 10 horas continuas entre jornadas.")
BULLET("Descanso semanal de al menos 35 horas continuas.")
BULLET("Registro fiel de jornada — hoy en Chile mediante dispositivo GPS homologado o libro de asistencia.")

P("Documentos que debes pedir", bold=True)
for d in [
    "Contratos de trabajo de los conductores con jornada explícita",
    "Libro de asistencia o registro GPS de los últimos 30 días",
    "Definición del régimen de trabajo (turno rotativo, excepcional, bisemanal)",
    "Resolución de la Dirección del Trabajo si hay régimen excepcional autorizado",
    "Política interna de fatiga y descansos",
]:
    BULLET(d)

P("Cómo lo controla el sistema", bold=True)
CODE("Tabla actividades_conductor     →  tipos: conduccion, espera, descanso, carga_descarga, pernocte")
CODE("Tabla registro_jornada_conductor →  resumen diario con alertas automáticas")
CODE("Función fn_verificar_fatiga_conductores  →  alerta 'fatiga_conductor'")

P("El GPS alimenta estas tablas automáticamente vía webhook en `fn_procesar_evento_gps`.")

CALLOUT(
    "Si un conductor excede las 5 horas continuas o las 88 horas mensuales de "
    "espera, el sistema dispara una alerta crítica. Si aun así sale a operar y "
    "ocurre un accidente, la responsabilidad civil y penal recae sobre la empresa."
)

# --------- RT, PC, SOAP ---------
H2("A.4  Documentos obligatorios de toda la flota")

TABLE(
    ["Documento", "Norma base", "Vigencia", "Dónde se obtiene"],
    [
        ["Revisión Técnica (RT)", "DS 156/1990 MTT", "Anual (vehículos pesados)", "Plantas de revisión técnica autorizadas"],
        ["Permiso de Circulación", "Ley de Tránsito 18.290", "Anual (marzo-mayo según región)", "Municipalidad"],
        ["SOAP", "Ley 18.490", "Anual", "Compañía de seguros"],
        ["Seguro de Responsabilidad Civil", "Exigencia contractual minera", "Anual", "Compañía de seguros"],
        ["FOPS / ROPS", "NCh / estándar ISO", "Según fabricante", "Inspección de estructura contra vuelcos y caída objetos"],
        ["Certificado de gancho", "Norma de equipos de izaje", "Anual", "Organismo certificador"],
    ],
)

P(
    "Todos estos se modelan en la tabla `certificaciones` del sistema. El campo "
    "`bloqueante` determina si impiden la disponibilidad del equipo."
)

doc.add_page_break()

# ============================================================
# BLOQUE B — SEGURIDAD LABORAL
# ============================================================
H1("Bloque B — Seguridad laboral y prevención de riesgos")

# --------- Ley 16.744 ---------
H2("B.1  Ley 16.744 — Seguro de accidentes del trabajo")

P("Qué regula", bold=True)
P(
    "Obliga al empleador a pagar cotización a una mutualidad (ACHS, Mutual de "
    "Seguridad, IST) y a prevenir, reportar y registrar accidentes y enfermedades "
    "profesionales. Es la base de TODA la prevención de riesgos en Chile."
)

P("Obligaciones del empleador", bold=True)
BULLET("Cotización mensual a mutualidad (0,95 % base + tasa adicional por riesgo).")
BULLET("Reglamento Interno de Higiene y Seguridad (RIHS) escrito, vigente y entregado a cada trabajador.")
BULLET("Derecho a saber (ODI — Obligación de Informar) sobre los riesgos del puesto, firmado por cada trabajador.")
BULLET("Comité Paritario de Higiene y Seguridad cuando la empresa tenga más de 25 trabajadores. Reuniones mensuales con actas.")
BULLET("Departamento de Prevención de Riesgos si la empresa tiene más de 100 trabajadores.")
BULLET("Reporte de accidente DENTRO de 24 horas a la mutualidad (DIAT — Declaración Individual de Accidente del Trabajo).")
BULLET("Reporte de enfermedad profesional (DIEP) cuando aplique.")
BULLET("Investigación formal de accidentes graves y fatales, con informe a la Dirección del Trabajo y SEREMI Salud.")
BULLET("Capacitación anual de los trabajadores en materias de seguridad.")

P("Documentos que debes pedir", bold=True)
for d in [
    "Certificado de adhesión a mutualidad al día",
    "Comprobantes de pago de cotizaciones últimos 6 meses",
    "RIHS vigente + constancias de entrega firmadas por cada trabajador",
    "Registros ODI por cada puesto de trabajo",
    "Actas del Comité Paritario últimos 12 meses",
    "Matriz de identificación de peligros y evaluación de riesgos (IPER)",
    "Plan anual de capacitaciones",
    "Estadística de accidentes últimos 12 meses (tasa, días perdidos, siniestralidad)",
]:
    BULLET(d)

P("Cómo lo controla el sistema", bold=True)
CODE("Tabla no_conformidades  →  tipo = 'incidente_seguridad'")
P(
    "Los incidentes se registran como no conformidades. La tasa de no conformidades "
    "impacta directamente el componente Calidad del OEE."
)

# --------- DS 132 ---------
H2("B.2  DS 132/2004 Minería — Reglamento de seguridad minera")

P("Qué regula", bold=True)
P(
    "Seguridad dentro de faenas mineras, tanto subterráneas como a rajo abierto. "
    "Aplica a Pillado porque la empresa opera dentro de sitios mineros (CMP Romeral, "
    "Cenizas Francke, Spence, etc). El cliente minero lo exige como condición "
    "contractual, y SERNAGEOMIN lo fiscaliza."
)

P("Obligaciones operacionales", bold=True)
BULLET("PTS — Procedimiento de Trabajo Seguro escrito por cada actividad crítica: carga/descarga de combustible, engrase en faena, cambio de neumático, operación de pluma, tránsito en rajo.")
BULLET("Inducción general de \"Hombre Nuevo\" antes de ingresar a la faena.")
BULLET("Inducciones específicas por área (rajo, planta, caminos internos).")
BULLET("Análisis de Riesgo del Trabajo (ART) o Análisis Seguro de Trabajo (AST) ANTES de iniciar cualquier trabajo en terreno.")
BULLET("Charla de 5 minutos diaria en faena.")
BULLET("Reporte inmediato a SERNAGEOMIN en caso de accidente grave o fatal.")
BULLET("Supervisor calificado en terreno (certificación minera).")

P("EPP mínimo exigido en faena minera:", bold=True)
BULLET("Casco con barbiquejo")
BULLET("Calzado de seguridad con punta de acero")
BULLET("Ropa de trabajo alta visibilidad (camisa manga larga + pantalón)")
BULLET("Antiparras o lentes de seguridad")
BULLET("Guantes según actividad")
BULLET("Protección auditiva cerca de equipos o plantas")
BULLET("Protección respiratoria cuando aplique")

P("Documentos que debes pedir", bold=True)
for d in [
    "PTS documentados de todas las actividades críticas (firmados por prevencionista y gerencia)",
    "Comprobantes de inducción de cada conductor/técnico que entre a faena",
    "Formatos de ART/AST en uso (última semana de respaldos)",
    "Matriz de EPP por cargo",
    "Actas de charlas de 5 minutos",
    "Copia del contrato de seguridad con el cliente minero",
    "Registro de incidentes reportados a SERNAGEOMIN últimos 12 meses",
]:
    BULLET(d)

# --------- Código Sanitario ---------
H2("B.3  Código Sanitario (DFL 725/1967) y fiscalización SEREMI Salud")

P("Qué regula", bold=True)
P(
    "Es el paraguas legal de la salud pública en Chile. De él cuelgan todos los "
    "reglamentos sanitarios: sustancias peligrosas (DS 43), residuos peligrosos "
    "(DS 148), agua potable, ruido, condiciones sanitarias básicas (DS 594)."
)

P("Aplicación directa en Pillado — DS 594/1999 Condiciones sanitarias y ambientales:", bold=True)
BULLET("Agua potable en faena para todos los trabajadores (mínimo 100 L por persona/día).")
BULLET("Servicios higiénicos (baños + duchas) proporcionales al número de trabajadores.")
BULLET("Comedor separado del área de trabajo, con condiciones de aseo.")
BULLET("Iluminación y ventilación mínimas en talleres.")
BULLET("Control de ruido (máx 85 dB promedio 8 horas).")
BULLET("Condiciones térmicas (DS 594 define límites de carga térmica).")
BULLET("Extintores operativos y señalización de emergencia.")

P("Fiscalización", bold=True)
P(
    "La SEREMI Salud regional (Coquimbo y Atacama para Pillado) puede inspeccionar "
    "sin aviso previo y aplicar multas directas. El Código Sanitario le da "
    "facultad de clausurar el lugar de trabajo si detecta riesgo grave."
)

doc.add_page_break()

# ============================================================
# BLOQUE C — MEDIO AMBIENTE
# ============================================================
H1("Bloque C — Medio ambiente (sustancias y residuos)")

P(
    "Este bloque es el más fiscalizado hoy en Chile y el que tiene multas más "
    "altas. Va a crecer en importancia con la llegada de RETC 2.0 y la trazabilidad "
    "electrónica obligatoria."
)

# --------- DS 43 ---------
H2("C.1  DS 43/2015 Salud — Almacenamiento de sustancias peligrosas (SUSPEL)")

P("Qué regula", bold=True)
P(
    "Condiciones mínimas de seguridad en bodegas donde se almacenen productos "
    "peligrosos: combustibles, aceites, solventes, ácidos, oxidantes, pinturas, "
    "pesticidas. En Pillado aplica al almacén de combustible, a la bodega de "
    "lubricantes del taller, y a cualquier estanque fijo de diesel."
)

P("Obligaciones de la bodega", bold=True)
BULLET("Autorización sanitaria de la SEREMI Salud regional antes de operar.")
BULLET("Autorización específica del Ministerio de Salud si almacena >30.000 kg.")
BULLET("Segregación por clase de peligro (no mezclar oxidantes con inflamables, etc.).")
BULLET("Sistema de contención secundaria (pretil o bandeja) con capacidad del 110 % del contenedor mayor.")
BULLET("Ducha de emergencia y lavaojos a distancia máxima de 10 metros del área de manipulación.")
BULLET("Kit anti-derrame con sorbentes, recipientes y EPP listo para usar.")
BULLET("Extintor(es) de tipo adecuado (ABC para inflamables, CO₂ en líquidos eléctricos).")
BULLET("Señalización: rombos NFPA, señalética de evacuación, prohibido fumar, obligación de EPP.")
BULLET("Hojas de Datos de Seguridad (HDS) vigentes de CADA producto, en idioma español y visibles.")
BULLET("Capacitación al personal que manipula sobre cada producto (registrada).")
BULLET("Plan de emergencia por derrame, incendio y fuga.")
BULLET("Inspección interna periódica (al menos mensual) con check-list registrado.")

P("Documentos que debes pedir", bold=True)
for d in [
    "Autorización sanitaria vigente de SEREMI Coquimbo y de SEREMI Atacama (si hay bodega en Calama)",
    "HDS (Hoja de Datos de Seguridad) de cada producto, versión vigente en español",
    "Plan de emergencia firmado por prevencionista",
    "Registro de capacitación del personal manipulador",
    "Actas de inspecciones mensuales",
    "Fotos del cumplimiento: ducha, lavaojos, contención, rotulado, extintores",
    "Comprobantes de calibración de balanzas si vende/despacha por peso",
]:
    BULLET(d)

P("Cómo lo controla el sistema", bold=True)
CODE("Tabla suspel_productos  →  catálogo con clase UN, pictogramas, HDS vigente")
CODE("Tabla suspel_bodegas    →  autorización sanitaria, equipamiento, inspecciones")
P("Vista `vw_prevencion_resumen` agrega el estado global al reporte diario.")

# --------- DS 148 ---------
H2("C.2  DS 148/2003 Salud — Manejo de residuos peligrosos (RESPEL)")

P("Qué regula", bold=True)
P(
    "Cómo debes clasificar, almacenar, transportar y disponer todo residuo que "
    "por su naturaleza (toxicidad, inflamabilidad, reactividad, corrosividad) sea "
    "peligroso. En Pillado se generan diariamente: aceite de motor usado, filtros "
    "usados, baterías, envases vacíos contaminados, trapos con hidrocarburos, "
    "neumáticos en desuso, lodos de trampas de grasa, anticongelantes."
)

P("Obligaciones", bold=True)
BULLET("Declarar la empresa como generador de RESPEL ante SEREMI Salud.")
BULLET("Plan de Manejo de Residuos Peligrosos aprobado si se genera > 12 ton/año totales O > 12 kg/año de residuo tóxico agudo.")
BULLET("Sitio de almacenamiento temporal autorizado, con contención, techumbre, segregación por tipo.")
BULLET("Almacenamiento máximo 6 meses (tóxicos agudos) o 12 meses (resto), bajo control SEREMI.")
BULLET("Libro de registro de generación y retiros, en papel o electrónico.")
BULLET("Entrega únicamente a empresa autorizada (Hidronor, Séché-Kadmar, Resin, Emeres, Proactiva, etc.).")
BULLET("Transportista autorizado por SEREMI para residuos peligrosos.")
BULLET("Certificado de disposición final emitido por el receptor por cada movimiento.")
BULLET("Declaración electrónica SIDREP dentro del plazo por cada movimiento generado y retirado.")

P("Qué es SIDREP y por qué importa", bold=True)
P(
    "SIDREP = Sistema de Declaración de Residuos Peligrosos. Es una plataforma "
    "electrónica del Ministerio del Medio Ambiente donde debes declarar cada vez "
    "que generas, trasladas o eliminas un residuo peligroso. Si no declaras, "
    "la multa va directamente al gerente general."
)

P("Documentos que debes pedir", bold=True)
for d in [
    "Plan de Manejo RESPEL aprobado por SEREMI Salud",
    "Usuario y contraseña SIDREP (Ventanilla Única RETC)",
    "Contratos vigentes con empresas receptoras autorizadas (Hidronor, Séché, Resin, etc.)",
    "Últimos 12 certificados de disposición final",
    "Libro de generación/retiros últimos 12 meses",
    "Clasificación de cada tipo de residuo (código Y, H, categoría peligrosidad)",
    "Permiso de almacenamiento temporal del sitio actual",
    "Resolución sanitaria del transportista contratado",
]:
    BULLET(d)

P("Cómo lo controla el sistema", bold=True)
CODE("Tabla respel_tipos                 →  catálogo de residuos generados")
CODE("Tabla respel_movimientos           →  libro de generación y retiros (con número SIDREP)")
CODE("Tabla respel_empresas_receptoras   →  receptores autorizados")

P(
    "En el reporte diario ves `respel_mes.generado_kg`, `respel_mes.retirado_kg`, "
    "y `respel_mes.pendientes_sidrep` (retiros sin declaración SIDREP hecha)."
)

# --------- NCh 382 ---------
H2("C.3  NCh 382 Of.2004 — Clasificación de sustancias peligrosas")

P("Qué es", bold=True)
P(
    "Norma chilena que adopta la clasificación de la ONU para sustancias "
    "peligrosas. Define las 9 clases ONU y los números UN de 4 dígitos que "
    "identifican cada producto. Es la base para el rotulado de camiones, etiquetas "
    "de envases y redacción de HDS."
)

P("Las 9 clases ONU que debes conocer:", bold=True)
TABLE(
    ["Clase", "Tipo", "Ejemplo típico en Pillado"],
    [
        ["1", "Explosivos", "No aplica — no se manejan en faena normal"],
        ["2", "Gases", "Acetileno, oxígeno (soldadura), LPG"],
        ["3", "Líquidos inflamables", "Diesel (UN 1202), Bencina (UN 1203), Kerosene (UN 1223)"],
        ["4", "Sólidos inflamables", "Magnesio, fósforo (menos habitual)"],
        ["5", "Oxidantes y peróxidos", "Nitrato de amonio (uso minero)"],
        ["6", "Tóxicos e infecciosos", "Algunos pesticidas"],
        ["7", "Radiactivos", "Solo con autorización específica"],
        ["8", "Corrosivos", "Ácido sulfúrico (UN 1830), hidróxido de sodio"],
        ["9", "Varios peligros", "Asbesto, baterías de litio usadas"],
    ],
)

P(
    "El campo `clase_un` en la tabla `suspel_productos` del sistema se llena "
    "con esta información. Ejemplo real: un camión que entrega diesel a CMP "
    "Romeral debe tener rombo rojo con llama y el número 1202 visible."
)

# --------- RETC ---------
H2("C.4  RETC — Registro de Emisiones y Transferencias de Contaminantes")

P("Qué es", bold=True)
P(
    "Sistema del Ministerio del Medio Ambiente que consolida en un solo lugar "
    "TODA la información ambiental que reportas. Obligación establecida en la "
    "Ley 20.417 y reglamentada por el DS 1/2013 MMA."
)

P("Sub-sistemas que se declaran por Ventanilla Única RETC:", bold=True)
BULLET("SIDREP — residuos peligrosos (generación, traslado, eliminación)")
BULLET("SINADER — residuos no peligrosos (escombros, chatarra, basura común)")
BULLET("DS 138 — emisiones atmosféricas de fuentes fijas (si tienes caldera, planta de asfalto, etc.)")
BULLET("Declaración anual de sustancias peligrosas almacenadas")
BULLET("RETC de emisiones al aire, agua, suelo")

P("Plazos habituales", bold=True)
P(
    "La declaración anual RETC del año calendario se presenta entre marzo y abril "
    "del año siguiente. SIDREP y SINADER son continuas (cada movimiento se "
    "declara dentro de los días hábiles posteriores)."
)

P("Documentos y accesos que debes pedir", bold=True)
for d in [
    "Usuario y contraseña de Ventanilla Única RETC (acceso de la empresa)",
    "Copia de la última declaración anual presentada",
    "Detalle por subsistema: SIDREP, SINADER, emisiones atmosféricas",
    "Respuesta a requerimientos o observaciones si las hubo",
]:
    BULLET(d)

# --------- SEIA breve ---------
H2("C.5  SEIA (Sistema de Evaluación de Impacto Ambiental) — cuándo aplica")

P(
    "El SEIA (Ley 19.300) aplica a proyectos de cierta envergadura. En Pillado "
    "NO ingresas proyectos al SEIA por la flota en sí, pero sí estás obligado "
    "a cumplir con las Resoluciones de Calificación Ambiental (RCA) de los "
    "clientes mineros donde operas."
)

P("Qué te puede pedir el cliente minero por su RCA:", bold=True)
BULLET("Solo combustible con cierto azufre máximo (ULSD).")
BULLET("Horario restringido de tránsito dentro del rajo.")
BULLET("Velocidad máxima y humectación de caminos internos.")
BULLET("Prohibición de lavado en sitios no autorizados.")
BULLET("Manejo específico de aguas de lavado de camiones.")
BULLET("Registro de cualquier derrame aunque sea menor.")

doc.add_page_break()

# ============================================================
# BLOQUE D — RESUMEN DE ALERTAS EN EL SISTEMA
# ============================================================
H1("Bloque D — Cómo el sistema conecta con cada norma")

P(
    "La función SQL `fn_ejecutar_verificaciones_normativas()` corre todas las "
    "verificaciones de una sola vez. Puedes ejecutarla en Supabase así:"
)
CODE("SELECT * FROM fn_ejecutar_verificaciones_normativas();")

P("Devuelve una tabla con cada verificación y el número de alertas generadas:")

TABLE(
    ["Verificación SQL", "Norma origen", "Dispara alerta", "Acción esperada"],
    [
        ["fn_verificar_antiguedad_flota", "DS 298", "antiguedad_vehiculo", "Dar de baja o reclasificar el vehículo"],
        ["fn_verificar_semep_conductores", "DS 298", "semep_vencido", "Renovar SEMEP antes que el conductor opere"],
        ["fn_verificar_fatiga_conductores", "Ley 21.561", "fatiga_conductor", "Dar descanso obligatorio inmediato"],
        ["fn_verificar_certificaciones_flota (rt)", "Ley Tránsito", "rt_por_vencer", "Agendar Revisión Técnica"],
        ["fn_verificar_certificaciones_flota (hermeticidad)", "DS 160", "hermeticidad_vencida", "Renovar prueba hidrostática"],
        ["fn_verificar_certificaciones_flota (sec)", "DS 160", "sec_no_vigente", "Renovar inscripción SEC / TC8"],
        ["fn_verificar_disponibilidad_vigente", "Control interno", "disponibilidad_vencida", "Ejecutar checklist de 55 ítems"],
        ["vw_prevencion_resumen (HDS)", "DS 43", "hds_por_revisar", "Actualizar HDS"],
        ["vw_prevencion_resumen (SIDREP)", "DS 148", "retiros_sin_sidrep", "Declarar movimiento en SIDREP"],
        ["vw_prevencion_resumen (bodegas)", "DS 43", "bodega_sin_autorizacion", "Gestionar autorización sanitaria"],
    ],
)

# ============================================================
# BLOQUE E — CHECKLIST MAESTRO
# ============================================================
doc.add_page_break()
H1("Bloque E — Checklist maestro de documentos")

H2("Por cada vehículo (55 en total)")
for c in [
    "Revisión Técnica vigente",
    "Permiso de Circulación vigente",
    "SOAP vigente",
    "Seguro de Responsabilidad Civil vigente",
    "TC8 / Aptitud Técnica (cisternas combustible)",
    "Certificado de Hermeticidad (cisternas)",
    "Inscripción SEC activa (cisternas)",
    "Certificado de gancho (camiones pluma)",
    "FOPS / ROPS cuando aplique",
    "Fotos del rotulado (cisternas cargas peligrosas)",
]:
    BULLET(c)

H2("Por cada conductor")
for c in [
    "Licencia profesional vigente (A2, A3, A4 o A5 según tipo)",
    "SEMEP vigente (anual para pesados)",
    "Curso de transporte de sustancias peligrosas",
    "ODI firmada del puesto",
    "Contrato con jornada definida (Ley 21.561)",
    "Registro de entrega de EPP",
    "Inducción de faena minera si aplica",
]:
    BULLET(c)

H2("Empresa Pillado")
for c in [
    "Adhesión a mutualidad al día",
    "RIHS vigente + firmas de entrega",
    "Actas Comité Paritario últimos 12 meses",
    "Matriz IPER actualizada",
    "Plan anual de capacitaciones",
    "Estadística accidentes últimos 12 meses",
]:
    BULLET(c)

H2("Bodegas y sitios de almacenamiento")
for c in [
    "Autorización sanitaria SEREMI Coquimbo (SUSPEL)",
    "Autorización sanitaria SEREMI Atacama (si corresponde)",
    "Plan de emergencia firmado por prevencionista",
    "HDS vigentes de cada producto almacenado",
    "Registro de capacitación del personal",
    "Inspecciones mensuales internas",
    "Permiso de almacenamiento temporal RESPEL",
]:
    BULLET(c)

H2("Gestión de residuos peligrosos")
for c in [
    "Plan de Manejo RESPEL aprobado por SEREMI",
    "Usuario y clave SIDREP",
    "Usuario y clave Ventanilla Única RETC",
    "Contratos vigentes con empresas receptoras",
    "Últimos 12 certificados de disposición final",
    "Libro de generación/retiros últimos 12 meses",
    "Declaración RETC año anterior presentada",
]:
    BULLET(c)

H2("Por cada contrato con cliente minero")
for c in [
    "Copia de la Resolución de Calificación Ambiental (RCA) del cliente",
    "PTS documentados de actividades críticas",
    "Comprobantes de inducción \"Hombre Nuevo\"",
    "Matriz de EPP aprobada por cliente",
    "Contactos de emergencia en faena",
    "Protocolos específicos del contrato",
]:
    BULLET(c)

# ============================================================
# CIERRE
# ============================================================
doc.add_paragraph()
H1("Consejo final de estudio")

P(
    "No memorices números de decreto — memoriza el mapa: \"para el vehículo DS 298 "
    "y DS 160; para el conductor SEMEP + Ley 21.561; para la operación DS 132 y "
    "Ley 16.744; para el ambiente DS 43 + DS 148 + RETC\". Con ese mapa ya puedes "
    "conversar con un fiscalizador o un cliente minero."
)
P(
    "Lo siguiente es, en el sistema, ejecutar `SELECT * FROM "
    "fn_ejecutar_verificaciones_normativas();` y entender cada alerta que aparece. "
    "Cada fila es una norma viva aplicándose a tu flota real."
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Documento de estudio — uso interno Pillado")
r.italic = True
r.font.size = Pt(9)

doc.save(OUT)
print(f"OK: {OUT} generado")
