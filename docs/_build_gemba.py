# -*- coding: utf-8 -*-
"""Genera el instructivo Word de puesta en marcha de Recorridos Gemba.

Sigue el estilo de los generadores de docs/onboarding-pillado (Calibri 11,
titulo como Title, margenes 2/2.2 cm, tablas Table Grid).
Uso: python docs/_build_gemba.py
"""
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL = RGBColor(0x0B, 0x2A, 0x4A)
GRIS = RGBColor(0x4B, 0x55, 0x63)
AMBAR = RGBColor(0xB4, 0x53, 0x09)
BASE = "https://pilladoiceo.netlify.app"

doc = Document()

# ── Estilos base ──
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(6)

for nivel, tam in ((1, 16), (2, 13), (3, 11)):
    h = doc.styles["Heading %d" % nivel]
    h.font.name = "Calibri"
    h.font.size = Pt(tam)
    h.font.color.rgb = AZUL
    h.font.bold = True

for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)


def link(par, texto, url):
    """Hipervínculo real (python-docx no lo trae de fábrica)."""
    r_id = par.part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(color)
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = texto
    run.append(t)
    h.append(run)
    par._p.append(h)
    return par


def p(texto="", bold=False, size=11, color=None, space_after=6, align=None, italic=False):
    par = doc.add_paragraph()
    if align:
        par.alignment = align
    par.paragraph_format.space_after = Pt(space_after)
    if texto:
        r = par.add_run(texto)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return par


def rich(partes, space_after=6, size=11):
    """Párrafo con tramos en negrita: [('texto', True/False), ...]"""
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    for texto, bold in partes:
        r = par.add_run(texto)
        r.bold = bold
        r.font.size = Pt(size)
    return par


def sombrear(celda, hexcolor):
    tcPr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def tabla(cabeceras, filas, anchos=None):
    t = doc.add_table(rows=1, cols=len(cabeceras))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, c in enumerate(cabeceras):
        hdr[i].text = ""
        par = hdr[i].paragraphs[0]
        par.paragraph_format.space_after = Pt(2)
        r = par.add_run(c)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        sombrear(hdr[i], "0B2A4A")
    for fila in filas:
        celdas = t.add_row().cells
        for i, v in enumerate(fila):
            celdas[i].text = ""
            par = celdas[i].paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            r = par.add_run(str(v))
            r.font.size = Pt(9.5)
    if anchos:
        for fila in t.rows:
            for i, w in enumerate(anchos):
                fila.cells[i].width = Cm(w)
    return t


def caja(titulo, cuerpo, fill="FFF7ED", color_titulo=AMBAR):
    """Aviso destacado, como recuadro de una sola celda."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    sombrear(c, fill)
    c.text = ""
    par = c.paragraphs[0]
    par.paragraph_format.space_after = Pt(2)
    r = par.add_run(titulo)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = color_titulo
    par2 = c.add_paragraph()
    par2.paragraph_format.space_after = Pt(2)
    r2 = par2.add_run(cuerpo)
    r2.font.size = Pt(9.5)
    return t


# ════════════════════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════════════════════
p("PILLADO Y CÍA. LTDA. · SICOM-ICEO", bold=True, size=9, color=GRIS, space_after=2)
doc.add_heading("Recorridos Gemba", level=0)
p("Puesta en marcha para Jefatura de Taller, Jefatura de Operaciones y Prevención de Riesgos",
  size=11.5, color=GRIS, space_after=14)

p("Los recorridos de terreno pasan del papel al sistema. Cada cargo tiene su propio checklist, "
  "se llena desde el celular mientras se camina, y lo que se encuentra queda con responsable y "
  "fecha de compromiso — no en una libreta.", space_after=10)

caja("Lo que cambia respecto de hoy",
     "Lo que se detecta ya no depende de que alguien se acuerde: queda registrado, con dueño y "
     "plazo, y se puede seguir hasta que se cierra. Y el avance del programa se mide solo: "
     "cuántos recorridos correspondían y cuántos se hicieron.")
p(space_after=10)

# ════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Correo para enviar a los involucrados", level=2)
p("Texto listo para copiar y pegar. El documento se adjunta.", italic=True, color=GRIS, size=9.5)

p(space_after=2)
rich([("Para: ", True), ("Jefe de Taller · Jefe de Operaciones · Prevención de Riesgos", False)], space_after=2)
rich([("Copia: ", True), ("Gerencia · Administración", False)], space_after=2)
rich([("Asunto: ", True), ("Recorridos Gemba en el sistema — ya disponible", False)], space_after=10)

p("Estimados:", space_after=8)

p("Desde hoy los recorridos de terreno se registran en SICOM-ICEO. Cada uno tiene su "
  "propio checklist según el cargo, se llena desde el celular caminando por el taller o la faena, "
  "y lo que se detecta queda como plan de acción con responsable y fecha.", space_after=8)

p("El objetivo es concreto: tener el taller impecable y asegurar la calidad de los equipos que "
  "entregamos. Por eso los recorridos de las jefaturas están orientados a la operación —que el "
  "plan se cumpla, que nada esté frenado, que el equipo salga sin observaciones— y el de "
  "Prevención a la seguridad. No son el mismo recorrido ni buscan lo mismo.", space_after=8)

p("Lo que le corresponde a cada uno:", bold=True, space_after=4)

tabla(
    ["Cargo", "Checklist", "Cuándo", "Cuánto toma"],
    [
        ["Jefe de Taller", "Operación y Calidad", "Todos los días", "10-15 min · 11-12 ítems"],
        ["Jefe de Operaciones", "Sistema y Cliente", "Cada dos semanas", "20-30 min · 23 ítems"],
        ["Prevención", "Caminata diaria de seguridad", "Todos los días", "10 min · 6 ítems"],
        ["Prevención", "Inspección planificada", "Una vez al mes", "45 min · 36 ítems"],
    ],
    anchos=[3.6, 4.8, 3.2, 4.6],
)
p(space_after=8)

p("El checklist del Jefe de Taller es corto porque rota: hay un núcleo fijo que se mira todos los "
  "días (si el plan se está ejecutando, qué está frenado, qué sale hoy al cliente) y un bloque que "
  "cambia según el día de la semana. En la semana se cubre el taller completo sin que ningún día "
  "tome más de 15 minutos.", space_after=8)

p("Links (se abren en el celular, con la cuenta de siempre):", bold=True, space_after=4)

par = doc.add_paragraph(style="List Bullet")
par.add_run("Recorridos Gemba: ").bold = True
link(par, BASE + "/dashboard/gemba", BASE + "/dashboard/gemba")

par = doc.add_paragraph(style="List Bullet")
par.add_run("Iniciar un recorrido: ").bold = True
link(par, BASE + "/dashboard/gemba/nuevo", BASE + "/dashboard/gemba/nuevo")

par = doc.add_paragraph(style="List Bullet")
par.add_run("Avance del programa: ").bold = True
link(par, BASE + "/dashboard/gemba/reporte", BASE + "/dashboard/gemba/reporte")

p(space_after=8)
p("Al entrar, el checklist que corresponde al cargo aparece ya seleccionado, y la portada avisa "
  "cuando toca recorrer. En el documento adjunto está el paso a paso.", space_after=8)

p("Cualquier duda o algo que no calce con la realidad del taller, me avisan y lo ajustamos: el "
  "contenido de los checklists se puede corregir.", space_after=8)

p("Saludos,", space_after=2)
p("Manuel Olivares", bold=True, space_after=0)
p("Operaciones · Pillado y Cía. Ltda.", size=9.5, color=GRIS, space_after=12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Cómo se hace un recorrido", level=2)
p("Cuatro pasos, desde el celular.", italic=True, color=GRIS, size=9.5, space_after=8)

pasos = [
    ("Entrar y abrir el módulo",
     "Ingresar con la cuenta de siempre y abrir Recorridos Gemba: Prevención de Riesgos entra por "
     "Prevención → Recorridos Gemba, y las jefaturas de Taller y de Operaciones por "
     "Operación → Recorridos de terreno. Si toca recorrer, "
     "aparece un aviso arriba con el botón para partir."),
    ("Iniciar el recorrido",
     "Botón «Nuevo recorrido». El checklist del cargo viene preseleccionado. Se elige el lugar "
     "(taller o faena), el sector y —si se quiere— un foco para ese día."),
    ("Marcar caminando",
     "Cada ítem se marca Cumple / No cumple / No aplica con el pulgar, y admite una observación. "
     "Al marcar «No cumple» se abre de inmediato el registro del hallazgo: qué se va a hacer, "
     "quién es el responsable y para cuándo."),
    ("Cerrar",
     "Cuando no quedan ítems pendientes se habilita «Cerrar recorrido». Ahí queda el registro con "
     "hora de término y el resultado."),
]
for i, (titulo, cuerpo) in enumerate(pasos, 1):
    rich([("Paso %d — %s. " % (i, titulo), True), (cuerpo, False)], space_after=8)

doc.add_heading("Dos reglas que conviene saber", level=3)

rich([("Un recorrido cerrado no se edita. ", True),
      ("Es el registro de lo que se vio ese día y sirve como respaldo ante una fiscalización o la "
       "mutual. Si hay que corregir un cierre por error, lo hace el administrador.", False)],
     space_after=6)

rich([("El plan de acción sigue vivo después del cierre. ", True),
      ("Una acción correctiva se cierra cuando efectivamente se hizo, aunque sea semanas después. "
       "Se marca desde el mismo recorrido o desde el plan de acción de la portada.", False)],
     space_after=10)

# ════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Qué mira cada checklist", level=2)

doc.add_heading("Jefe de Taller — Operación y calidad del equipo (diario)", level=3)
p("Núcleo fijo, todos los días:", bold=True, space_after=3)
for t in [
    "El plan de hoy está publicado y cada técnico sabe en qué OT parte.",
    "Las OT marcadas en ejecución coinciden con lo que se hace en los boxes.",
    "Los equipos detenidos esperando repuesto están identificados y con su vale emitido.",
    "Ningún vale pendiente de despacho lleva más de 24 h sin escalar.",
    "Ningún equipo estacionado sin OT abierta ni fecha de salida comprometida.",
    "Se le preguntó a un técnico qué le impide avanzar y quedó registrado.",
    "No se toleró ninguna condición o acto de riesgo evidente.",
]:
    doc.add_paragraph(t, style="List Bullet")

p("Bloque que rota por día de la semana:", bold=True, space_after=3)
tabla(
    ["Día", "Qué se mira a fondo"],
    [
        ["Lunes", "Plan y carga de la semana"],
        ["Martes", "Calidad de lo ejecutado (pauta completa, evidencia, mediciones con instrumento)"],
        ["Miércoles", "El equipo que sale al cliente (documentación, limpieza, certificados, retrabajo)"],
        ["Jueves", "El taller como herramienta (orden, repuestos identificados, zonas demarcadas)"],
        ["Viernes", "Personas, estándar y cierre de semana"],
    ],
    anchos=[3.0, 13.2],
)
p(space_after=10)

doc.add_heading("Jefe de Operaciones — Sistema y cliente (quincenal)", level=3)
for t in [
    "El compromiso con el cliente se cumple: equipos y servicios comprometidos, reclamos con cierre.",
    "Calidad de lo que sale del taller: se revisan en terreno 2 o 3 equipos entregados.",
    "El taller funciona como sistema: plan, preventivas vencidas, NC en plazo, repuestos.",
    "Personas y obstáculos: se conversa con técnicos y operadores, y lo que piden tiene respuesta.",
    "En faena: estándar del mandante, registro el mismo día, documentación al día, coordinación.",
]:
    doc.add_paragraph(t, style="List Bullet")
p(space_after=8)

doc.add_heading("Prevención — Seguridad (diario + mensual)", level=3)
rich([("Caminata diaria (6 ítems). ", True),
      ("Observar una tarea completa de principio a fin, conversar con quien la ejecuta y corregir "
       "en el momento. Es presencia, no auditoría.", False)], space_after=6)
rich([("Inspección planificada mensual (36 ítems). ", True),
      ("EPP, orden y aseo, máquinas, riesgo eléctrico, sustancias peligrosas, emergencias, "
       "ergonomía y comportamientos. Va entera: es el respaldo formal del programa preventivo.", False)],
     space_after=10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Cómo se mide el avance", level=2)

par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(8)
par.add_run("El reporte de avance (").font.size = Pt(10.5)
link(par, BASE + "/dashboard/gemba/reporte", BASE + "/dashboard/gemba/reporte")
r = par.add_run(") muestra, por mes:")
r.font.size = Pt(11)

tabla(
    ["Indicador", "Qué responde"],
    [
        ["Avance del programa", "Cuántos recorridos correspondían y cuántos se hicieron, por checklist"],
        ["Plan de acción", "Abiertos, vencidos, cerrados y qué porcentaje llegó dentro del plazo"],
        ["Lo que más falla", "Los ítems con más «no cumple» en 90 días"],
        ["Quién recorrió", "Recorridos por persona y hace cuántos días fue el último"],
    ],
    anchos=[4.6, 11.6],
)
p(space_after=8)

caja("Por qué hay dos porcentajes y no uno",
     "El «avance del programa» dice si se sale a terreno; el «% de cumplimiento del checklist» dice "
     "qué tan bien salió. Van separados a propósito: el segundo sube marcando «cumple» sin mirar, "
     "el primero no se puede falsear sin salir a caminar. Si hay que elegir un solo número para la "
     "reunión, es el del programa.",
     fill="EFF6FF", color_titulo=AZUL)
p(space_after=10)

# ════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Accesos", level=2)
p("El módulo ya está publicado y los tres links de este documento están operativos. "
  "Cada uno entra con su propia cuenta:", space_after=6)

tabla(
    ["Quién", "Cuenta", "Qué ve"],
    [
        ["Jefe de Taller", "La de siempre", "Su checklist diario, ya habilitado"],
        ["Jefe de Operaciones", "La de siempre", "Su checklist quincenal, ya habilitado"],
        ["Prevención — Anyulin Cortes", "prevcoq@pillado.cl", "Caminata diaria + inspección mensual"],
    ],
    anchos=[5.2, 5.0, 6.0],
)
p(space_after=8)

rich([("Prevención tiene cuenta nueva. ", True),
      ("La contraseña inicial se entrega por separado, no en este documento. Conviene cambiarla "
       "en el primer ingreso.", False)], space_after=6)

rich([("Al entrar, cada uno ve solo el checklist de su cargo preseleccionado ", True),
      ("y, si le toca recorrer, un aviso en la portada del módulo con el botón para partir.", False)],
     space_after=10)

caja("Si algo del checklist no calza con la realidad del taller",
     "El contenido de cada checklist es editable: se pueden cambiar ítems, agregar o sacar, sin "
     "rehacer el sistema. Conviene ajustarlo con lo que digan los propios jefes en las primeras "
     "dos semanas de uso — un checklist que no calza con el terreno se empieza a llenar sin mirar.")

p(space_after=14)
p("Documento generado para la puesta en marcha del módulo de Recorridos Gemba · SICOM-ICEO",
  size=8.5, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

import os
salida = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                      "Recorridos_Gemba_Puesta_en_Marcha.docx")
doc.save(salida)
print("OK ->", salida)
